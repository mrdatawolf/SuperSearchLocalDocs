"""Document indexer - Extracts text from documents and stores in SQLite database"""

import os
import sqlite3
from pathlib import Path
from datetime import datetime
import traceback
import multiprocessing
from functools import partial

# Document parsing libraries
try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

from config import DOCUMENT_PATH, DATABASE_PATH, SUPPORTED_EXTENSIONS

# Parallel processing configuration
ENABLE_PARALLEL = True  # Enable parallel processing by default
MAX_WORKERS = None  # None = auto-detect (75% of CPU cores)


class DocumentIndexer:
    def __init__(self, document_path=None, db_path=None):
        """
        Initialize the document indexer

        Args:
            document_path: Path to folder containing documents to index (uses config default if None)
            db_path: Path to SQLite database file (uses config default if None)
        """
        self.db_path = db_path or DATABASE_PATH
        self.document_path = document_path or DOCUMENT_PATH
        print(f"\n{'='*80}")
        print(f"INDEXER CONFIGURATION:")
        print(f"{'='*80}")
        print(f"Database path: {self.db_path}")
        print(f"Absolute database path: {Path(self.db_path).absolute()}")
        print(f"Database exists: {os.path.exists(self.db_path)}")
        print(f"Document path: {self.document_path}")
        print(f"{'='*80}\n")

        # Ensure database directory exists
        Path(self.db_path).parent.mkdir(parents=True, exist_ok=True)

        self.init_database()

    def init_database(self):
        """Initialize SQLite database with FTS5 for full-text search"""
        conn = sqlite3.connect(self.db_path, timeout=30.0)
        cursor = conn.cursor()

        # Create documents table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_path TEXT UNIQUE NOT NULL,
                file_name TEXT NOT NULL,
                folder_path TEXT,
                file_type TEXT NOT NULL,
                file_size INTEGER,
                modified_date TEXT,
                indexed_date TEXT NOT NULL,
                content TEXT
            )
        ''')

        # Create FTS5 virtual table for full-text search with multiple searchable columns
        cursor.execute('''
            CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(
                file_name,
                folder_path,
                content,
                content='documents',
                content_rowid='id'
            )
        ''')

        # Create word_counts table for pre-calculated popular words
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS word_counts (
                word TEXT PRIMARY KEY,
                count INTEGER NOT NULL DEFAULT 1
            )
        ''')

        # Create index on count for faster top-N queries
        cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_word_counts_count ON word_counts(count DESC)
        ''')

        # Create triggers to keep FTS table in sync
        cursor.execute('''
            CREATE TRIGGER IF NOT EXISTS documents_ai AFTER INSERT ON documents BEGIN
                INSERT INTO documents_fts(rowid, file_name, folder_path, content)
                VALUES (new.id, new.file_name, new.folder_path, new.content);
            END;
        ''')

        cursor.execute('''
            CREATE TRIGGER IF NOT EXISTS documents_au AFTER UPDATE ON documents BEGIN
                UPDATE documents_fts
                SET file_name = new.file_name,
                    folder_path = new.folder_path,
                    content = new.content
                WHERE rowid = new.id;
            END;
        ''')

        cursor.execute('''
            CREATE TRIGGER IF NOT EXISTS documents_ad AFTER DELETE ON documents BEGIN
                DELETE FROM documents_fts WHERE rowid = old.id;
            END;
        ''')

        conn.commit()
        conn.close()
        print(f"Database initialized at {self.db_path}")

    def update_word_counts(self, content):
        """Extract words from content and update word_counts table"""
        if not content or len(content) < 20:
            return

        import re

        # Stop words to exclude
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                     'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'be', 'been',
                     'this', 'that', 'these', 'those', 'it', 'its', 'can', 'will', 'would',
                     'sheet', 'none', 'true', 'false', 'openpyxl', 'not', 'installed',
                     'error', 'reading', 'has', 'have', 'had', 'do', 'does', 'did', 'you',
                     'your', 'we', 'our', 'they', 'their', 'he', 'she', 'his', 'her', 'if',
                     'then', 'than', 'so', 'what', 'when', 'where', 'who', 'which', 'how',
                     'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some',
                     'such', 'no', 'nor', 'only', 'own', 'same', 'than', 'too', 'very', 'nan'}

        # Extract words (3+ chars, letters only)
        words = re.findall(r'\b[a-z]{3,}\b', content.lower())

        # Count words (excluding stop words)
        word_freq = {}
        for word in words:
            if word not in stop_words:
                word_freq[word] = word_freq.get(word, 0) + 1

        # Update database with word counts
        if word_freq:
            conn = sqlite3.connect(self.db_path, timeout=30.0)
            cursor = conn.cursor()

            try:
                for word, count in word_freq.items():
                    # Insert or update word count
                    cursor.execute('''
                        INSERT INTO word_counts (word, count)
                        VALUES (?, ?)
                        ON CONFLICT(word) DO UPDATE SET count = count + ?
                    ''', (word, count, count))

                conn.commit()
            except Exception as e:
                print(f"Warning: Failed to update word counts: {e}")
            finally:
                conn.close()

    def rebuild_word_counts(self):
        """Rebuild word_counts table from all existing documents - useful for existing databases"""
        import re

        print("\nRebuilding word counts from existing documents...")

        # Stop words to exclude
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                     'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'be', 'been',
                     'this', 'that', 'these', 'those', 'it', 'its', 'can', 'will', 'would',
                     'sheet', 'none', 'true', 'false', 'openpyxl', 'not', 'installed',
                     'error', 'reading', 'has', 'have', 'had', 'do', 'does', 'did', 'you',
                     'your', 'we', 'our', 'they', 'their', 'he', 'she', 'his', 'her', 'if',
                     'then', 'than', 'so', 'what', 'when', 'where', 'who', 'which', 'how',
                     'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some',
                     'such', 'no', 'nor', 'only', 'own', 'same', 'than', 'too', 'very', 'nan'}

        conn = sqlite3.connect(self.db_path, timeout=30.0)
        cursor = conn.cursor()

        try:
            # Get total document count
            cursor.execute('SELECT COUNT(*) FROM documents WHERE content IS NOT NULL AND LENGTH(content) > 20')
            total_docs = cursor.fetchone()[0]
            print(f"Processing {total_docs} documents...")

            # Build word counts in memory
            all_word_counts = {}

            # Process documents in batches using ROWID cursor (much faster than OFFSET)
            batch_size = 1000  # Increased batch size for better performance
            last_rowid = 0
            processed = 0

            while True:
                cursor.execute('''
                    SELECT rowid, content FROM documents
                    WHERE rowid > ?
                      AND content IS NOT NULL
                      AND LENGTH(content) > 20
                    ORDER BY rowid
                    LIMIT ?
                ''', (last_rowid, batch_size))

                rows = cursor.fetchall()
                if not rows:
                    break

                processed += len(rows)
                print(f"Processing documents {processed - len(rows) + 1} to {processed}...")

                for rowid, content in rows:
                    # Extract words (3+ chars, letters only)
                    words = re.findall(r'\b[a-z]{3,}\b', content.lower())

                    # Count words (excluding stop words)
                    for word in words:
                        if word not in stop_words:
                            all_word_counts[word] = all_word_counts.get(word, 0) + 1

                    last_rowid = rowid

            print(f"Found {len(all_word_counts)} unique words")
            print("Clearing old word counts and inserting new data...")

            # Clear existing word counts and insert new ones in a single transaction
            cursor.execute('DELETE FROM word_counts')

            # Insert all word counts
            for word, count in all_word_counts.items():
                cursor.execute('''
                    INSERT INTO word_counts (word, count)
                    VALUES (?, ?)
                ''', (word, count))

            conn.commit()
            print("[OK] Word counts rebuilt successfully!")

        except Exception as e:
            print(f"Error rebuilding word counts: {e}")
            conn.rollback()
        finally:
            conn.close()

    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX files"""
        if Document is None:
            return "[python-docx not installed]"

        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            return f"[Error reading DOCX: {str(e)}]"

    def extract_text_from_csv(self, file_path):
        """Extract text from CSV files"""
        if pd is None:
            return "[pandas not installed]"

        try:
            df = pd.read_csv(file_path, encoding='utf-8', on_bad_lines='skip')
            return df.to_string()
        except Exception as e:
            try:
                df = pd.read_csv(file_path, encoding='latin-1', on_bad_lines='skip')
                return df.to_string()
            except:
                return f"[Error reading CSV: {str(e)}]"

    def extract_text_from_xlsx(self, file_path):
        """Extract text from XLSX files"""
        if load_workbook is None:
            return "[openpyxl not installed]"

        try:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            text = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text.append(f"Sheet: {sheet_name}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        text.append(row_text)
            return '\n'.join(text)
        except Exception as e:
            return f"[Error reading XLSX: {str(e)}]"

    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF files"""
        if PdfReader is None:
            return "[PyPDF2 not installed]"

        try:
            reader = PdfReader(file_path)
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return '\n'.join(text)
        except Exception as e:
            return f"[Error reading PDF: {str(e)}]"

    def extract_text_from_image(self, file_path):
        """Extract text from images using OCR"""
        if Image is None or pytesseract is None:
            return "[Pillow or pytesseract not installed]"

        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            return text if text.strip() else "[No text found in image]"
        except Exception as e:
            return f"[Error reading image: {str(e)}]"

    def extract_text_from_ps1(self, file_path):
        """Extract text from PowerShell script files"""
        try:
            # Try UTF-8 first (most common for modern scripts)
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                # Fallback to UTF-16 (sometimes used by Windows)
                with open(file_path, 'r', encoding='utf-16') as f:
                    return f.read()
            except UnicodeDecodeError:
                try:
                    # Final fallback to latin-1 (always succeeds)
                    with open(file_path, 'r', encoding='latin-1') as f:
                        return f.read()
                except Exception as e:
                    return f"[Error reading PowerShell script: {str(e)}]"
        except Exception as e:
            return f"[Error reading PowerShell script: {str(e)}]"

    def extract_text_from_txt(self, file_path):
        """Extract text from plain text files"""
        try:
            # Try UTF-8 first (most common)
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                return f"[Error reading text file: {str(e)}]"
        except Exception as e:
            return f"[Error reading text file: {str(e)}]"

    def extract_text_from_xls(self, file_path):
        """Extract text from legacy Excel files (.xls)"""
        if pd is None:
            return "[pandas not installed]"

        try:
            # pandas can read .xls files using xlrd
            df_dict = pd.read_excel(file_path, sheet_name=None, engine='xlrd')
            text = []
            for sheet_name, df in df_dict.items():
                text.append(f"Sheet: {sheet_name}")
                text.append(df.to_string())
            return '\n'.join(text)
        except Exception as e:
            return f"[Error reading XLS: {str(e)}. Note: xlrd package may be required for .xls files]"

    def extract_text(self, file_path):
        """Extract text based on file extension"""
        ext = Path(file_path).suffix.lower()

        if ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.csv':
            return self.extract_text_from_csv(file_path)
        elif ext == '.xlsx':
            return self.extract_text_from_xlsx(file_path)
        elif ext == '.xls':
            return self.extract_text_from_xls(file_path)
        elif ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.txt':
            return self.extract_text_from_txt(file_path)
        elif ext == '.ps1':
            return self.extract_text_from_ps1(file_path)
        elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif']:
            return self.extract_text_from_image(file_path)
        else:
            return "[Unsupported file type]"

    @staticmethod
    def extract_document_data(file_path, document_path):
        """
        Worker function for parallel processing - extracts document metadata and text
        This runs in a separate process, so it doesn't access self or the database

        Returns: (success, file_path, metadata_dict, content) or (success, file_path, error_message, None)
        """
        try:
            path = Path(file_path)

            # Get file metadata
            file_name = path.name
            file_type = SUPPORTED_EXTENSIONS.get(path.suffix.lower(), 'Unknown')
            file_size = path.stat().st_size
            modified_date = datetime.fromtimestamp(path.stat().st_mtime).isoformat()
            indexed_date = datetime.now().isoformat()

            # Extract folder path
            try:
                folder_path = str(path.parent.relative_to(Path(document_path)))
                folder_path = folder_path.replace('\\', ' ').replace('/', ' ')
            except ValueError:
                folder_path = str(path.parent)

            # Extract text content - this is the CPU-intensive part
            ext = path.suffix.lower()

            # Use static methods to extract text (since we can't access self in worker)
            if ext == '.docx':
                content = DocumentIndexer._static_extract_docx(file_path)
            elif ext == '.csv':
                content = DocumentIndexer._static_extract_csv(file_path)
            elif ext == '.xlsx':
                content = DocumentIndexer._static_extract_xlsx(file_path)
            elif ext == '.xls':
                content = DocumentIndexer._static_extract_xls(file_path)
            elif ext == '.pdf':
                content = DocumentIndexer._static_extract_pdf(file_path)
            elif ext == '.txt':
                content = DocumentIndexer._static_extract_txt(file_path)
            elif ext == '.ps1':
                content = DocumentIndexer._static_extract_ps1(file_path)
            elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif']:
                content = DocumentIndexer._static_extract_image(file_path)
            else:
                content = "[Unsupported file type]"

            # Package metadata
            metadata = {
                'file_name': file_name,
                'folder_path': folder_path,
                'file_type': file_type,
                'file_size': file_size,
                'modified_date': modified_date,
                'indexed_date': indexed_date
            }

            return (True, str(file_path), metadata, content)

        except Exception as e:
            return (False, str(file_path), str(e), None)

    # Static extraction methods for workers
    @staticmethod
    def _static_extract_docx(file_path):
        """Static version of DOCX extraction for worker processes"""
        if Document is None:
            return "[python-docx not installed]"
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            return f"[Error reading DOCX: {str(e)}]"

    @staticmethod
    def _static_extract_csv(file_path):
        """Static version of CSV extraction for worker processes"""
        if pd is None:
            return "[pandas not installed]"
        try:
            df = pd.read_csv(file_path, encoding='utf-8', on_bad_lines='skip')
            return df.to_string()
        except Exception:
            try:
                df = pd.read_csv(file_path, encoding='latin-1', on_bad_lines='skip')
                return df.to_string()
            except Exception as e:
                return f"[Error reading CSV: {str(e)}]"

    @staticmethod
    def _static_extract_xlsx(file_path):
        """Static version of XLSX extraction for worker processes"""
        if load_workbook is None:
            return "[openpyxl not installed]"
        try:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            text = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text.append(f"Sheet: {sheet_name}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        text.append(row_text)
            return '\n'.join(text)
        except Exception as e:
            return f"[Error reading XLSX: {str(e)}]"

    @staticmethod
    def _static_extract_pdf(file_path):
        """Static version of PDF extraction for worker processes"""
        if PdfReader is None:
            return "[PyPDF2 not installed]"
        try:
            reader = PdfReader(file_path)
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return '\n'.join(text)
        except Exception as e:
            return f"[Error reading PDF: {str(e)}]"

    @staticmethod
    def _static_extract_image(file_path):
        """Static version of image extraction for worker processes"""
        if Image is None or pytesseract is None:
            return "[Pillow or pytesseract not installed]"
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            return text if text.strip() else "[No text found in image]"
        except Exception as e:
            return f"[Error reading image: {str(e)}]"

    @staticmethod
    def _static_extract_ps1(file_path):
        """Static version of PowerShell script extraction for worker processes"""
        try:
            # Try UTF-8 first (most common for modern scripts)
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                # Fallback to UTF-16 (sometimes used by Windows)
                with open(file_path, 'r', encoding='utf-16') as f:
                    return f.read()
            except UnicodeDecodeError:
                try:
                    # Final fallback to latin-1 (always succeeds)
                    with open(file_path, 'r', encoding='latin-1') as f:
                        return f.read()
                except Exception as e:
                    return f"[Error reading PowerShell script: {str(e)}]"
        except Exception as e:
            return f"[Error reading PowerShell script: {str(e)}]"

    @staticmethod
    def _static_extract_txt(file_path):
        """Static version of text file extraction for worker processes"""
        try:
            # Try UTF-8 first (most common)
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                return f"[Error reading text file: {str(e)}]"
        except Exception as e:
            return f"[Error reading text file: {str(e)}]"

    @staticmethod
    def _static_extract_xls(file_path):
        """Static version of legacy Excel extraction for worker processes"""
        if pd is None:
            return "[pandas not installed]"

        try:
            # pandas can read .xls files using xlrd
            df_dict = pd.read_excel(file_path, sheet_name=None, engine='xlrd')
            text = []
            for sheet_name, df in df_dict.items():
                text.append(f"Sheet: {sheet_name}")
                text.append(df.to_string())
            return '\n'.join(text)
        except Exception as e:
            return f"[Error reading XLS: {str(e)}. Note: xlrd package may be required for .xls files]"

    def index_document(self, file_path):
        """Index a single document"""
        try:
            path = Path(file_path)

            # Get file metadata
            file_name = path.name
            file_type = SUPPORTED_EXTENSIONS.get(path.suffix.lower(), 'Unknown')
            file_size = path.stat().st_size
            modified_date = datetime.fromtimestamp(path.stat().st_mtime).isoformat()
            indexed_date = datetime.now().isoformat()

            # Extract folder path (relative to document root, including all subfolders)
            try:
                folder_path = str(path.parent.relative_to(Path(self.document_path)))
                # Replace path separators with spaces for better searchability
                folder_path = folder_path.replace('\\', ' ').replace('/', ' ')
            except ValueError:
                # If path is not relative to document_path, use the full parent path
                folder_path = str(path.parent)

            # Extract text content
            content = self.extract_text(file_path)

            # Store in database with retry logic for locked database
            max_retries = 3
            retry_delay = 1

            for attempt in range(max_retries):
                try:
                    conn = sqlite3.connect(self.db_path, timeout=30.0)
                    cursor = conn.cursor()

                    # Check if document already exists
                    cursor.execute('SELECT id FROM documents WHERE file_path = ?', (str(file_path),))
                    existing = cursor.fetchone()

                    if existing:
                        # Update existing document (trigger will update FTS)
                        doc_id = existing[0]
                        cursor.execute('''
                            UPDATE documents
                            SET file_name=?, folder_path=?, file_type=?, file_size=?, modified_date=?, indexed_date=?, content=?
                            WHERE id=?
                        ''', (file_name, folder_path, file_type, file_size, modified_date, indexed_date, content, doc_id))
                    else:
                        # Insert new document (trigger will insert into FTS)
                        cursor.execute('''
                            INSERT INTO documents (file_path, file_name, folder_path, file_type, file_size, modified_date, indexed_date, content)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                        ''', (str(file_path), file_name, folder_path, file_type, file_size, modified_date, indexed_date, content))

                    conn.commit()
                    conn.close()

                    # Update word counts for this document's content
                    self.update_word_counts(content)

                    return True, f"Indexed: {file_name}"

                except sqlite3.OperationalError as e:
                    if "locked" in str(e).lower() and attempt < max_retries - 1:
                        # Database is locked, retry after delay
                        import time
                        time.sleep(retry_delay)
                        continue
                    else:
                        raise

        except Exception as e:
            return False, f"Error indexing {file_path}: {str(e)}"

    def write_document_to_db(self, file_path, metadata, content):
        """
        Write extracted document data to database
        Used by parallel processing to write results from workers
        """
        try:
            # Debug: Show database path
            # print(f"DEBUG: Writing to database: {self.db_path}")

            conn = sqlite3.connect(self.db_path, timeout=30.0)
            cursor = conn.cursor()

            # Check if document already exists
            cursor.execute('SELECT id FROM documents WHERE file_path = ?', (file_path,))
            existing = cursor.fetchone()

            if existing:
                # Update existing document
                doc_id = existing[0]
                cursor.execute('''
                    UPDATE documents
                    SET file_name=?, folder_path=?, file_type=?, file_size=?, modified_date=?, indexed_date=?, content=?
                    WHERE id=?
                ''', (metadata['file_name'], metadata['folder_path'], metadata['file_type'],
                      metadata['file_size'], metadata['modified_date'], metadata['indexed_date'],
                      content, doc_id))
            else:
                # Insert new document
                cursor.execute('''
                    INSERT INTO documents (file_path, file_name, folder_path, file_type, file_size, modified_date, indexed_date, content)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''', (file_path, metadata['file_name'], metadata['folder_path'], metadata['file_type'],
                      metadata['file_size'], metadata['modified_date'], metadata['indexed_date'], content))

            conn.commit()
            conn.close()

            # Update word counts for this document's content
            self.update_word_counts(content)

            return True, f"Indexed: {metadata['file_name']}"

        except Exception as e:
            import traceback
            error_detail = traceback.format_exc()
            print(f"Database write error details:\n{error_detail}")
            return False, f"Error writing {Path(file_path).name} to database: {str(e)}"

    def scan_and_index(self, use_parallel=None, max_workers=None, progress_callback=None):
        """
        Scan document folder and index all supported files

        Args:
            use_parallel: Enable parallel processing (default: ENABLE_PARALLEL constant)
            max_workers: Number of worker processes (default: 75% of CPU cores)
            progress_callback: Optional callback function(indexed, total, message) for progress updates
        """
        print(f"Scanning documents in: {self.document_path}")
        print("-" * 80)

        if not os.path.exists(self.document_path):
            print(f"ERROR: Path does not exist: {self.document_path}")
            return

        # Determine if we should use parallel processing
        if use_parallel is None:
            use_parallel = ENABLE_PARALLEL

        # Determine worker count
        if max_workers is None:
            if MAX_WORKERS is None:
                # Auto-detect: use 75% of CPU cores, minimum 1
                cpu_count = os.cpu_count() or 1
                max_workers = max(1, int(cpu_count * 0.75))
            else:
                max_workers = MAX_WORKERS

        # Build list of files to index
        files_to_index = []
        skipped_extensions = {}
        total_scanned = 0

        for root, dirs, files in os.walk(self.document_path):
            for file in files:
                total_scanned += 1
                ext = Path(file).suffix.lower()
                if ext in SUPPORTED_EXTENSIONS:
                    file_path = os.path.join(root, file)
                    files_to_index.append(file_path)
                else:
                    # Track skipped extensions
                    if ext:
                        skipped_extensions[ext] = skipped_extensions.get(ext, 0) + 1

        total_files = len(files_to_index)

        # Show scanning summary
        print(f"Scanned {total_scanned} total files")
        print(f"Found {total_files} supported documents to index")
        if skipped_extensions:
            print(f"Skipped {sum(skipped_extensions.values())} files with unsupported extensions:")
            # Show top 10 skipped extensions
            sorted_skipped = sorted(skipped_extensions.items(), key=lambda x: x[1], reverse=True)
            for ext, count in sorted_skipped[:10]:
                ext_display = ext if ext else "(no extension)"
                print(f"  {ext_display}: {count} files")

        if total_files == 0:
            print("No supported documents found to index")
            print(f"Supported extensions: {', '.join(SUPPORTED_EXTENSIONS.keys())}")
            return

        # Decide mode based on file count and parallel setting
        if use_parallel and total_files > 5:  # Only use parallel for more than 5 files
            print(f"Parallel processing enabled: using {max_workers} of {os.cpu_count()} CPU cores")
            print(f"Found {total_files} documents to process")
            print("-" * 80)
            indexed_count, error_count = self._scan_parallel(files_to_index, max_workers, progress_callback)
        else:
            if use_parallel and total_files <= 5:
                print(f"Found {total_files} documents - using sequential mode (parallel not efficient for small batches)")
            else:
                print(f"Sequential mode - Found {total_files} documents to process")
            print("-" * 80)
            indexed_count, error_count = self._scan_sequential(files_to_index, progress_callback)

        print("-" * 80)
        print(f"Indexing complete!")
        print(f"Successfully indexed: {indexed_count} documents")
        print(f"Errors: {error_count}")

        # Verify database contents
        try:
            conn = sqlite3.connect(self.db_path, timeout=30.0)
            cursor = conn.cursor()
            cursor.execute('SELECT COUNT(*) FROM documents')
            db_count = cursor.fetchone()[0]
            conn.close()
            print(f"Database now contains: {db_count} total documents")

            if indexed_count > 0 and db_count == 0:
                print("⚠️  WARNING: Files were processed but database is empty!")
                print(f"   Database path: {self.db_path}")
                print(f"   Database exists: {os.path.exists(self.db_path)}")
        except Exception as e:
            print(f"Could not verify database: {e}")

        return indexed_count, error_count

    def _scan_sequential(self, files_to_index, progress_callback=None):
        """Sequential indexing (original method)"""
        indexed_count = 0
        error_count = 0
        total = len(files_to_index)

        for file_path in files_to_index:
            success, message = self.index_document(file_path)

            if success:
                indexed_count += 1
                print(f"✓ {message}")
            else:
                error_count += 1
                print(f"✗ {message}")

            if progress_callback:
                progress_callback(indexed_count + error_count, total, message)

        return indexed_count, error_count

    def _scan_parallel(self, files_to_index, max_workers, progress_callback=None):
        """Parallel indexing using multiprocessing"""
        indexed_count = 0
        error_count = 0
        total = len(files_to_index)

        # Create worker function with document_path bound
        worker_func = partial(DocumentIndexer.extract_document_data, document_path=self.document_path)

        try:
            # Create process pool
            with multiprocessing.Pool(processes=max_workers) as pool:
                # Process files in parallel
                # Use imap_unordered for better performance (results come as workers finish)
                results = pool.imap_unordered(worker_func, files_to_index, chunksize=2)

                # Process results from workers
                for success, file_path, data, content in results:
                    message = ""
                    if success:
                        # Write to database (main process only)
                        db_success, message = self.write_document_to_db(file_path, data, content)

                        if db_success:
                            indexed_count += 1
                            print(f"✓ {message}")
                        else:
                            error_count += 1
                            print(f"✗ {message}")
                    else:
                        # data contains error message in this case
                        error_count += 1
                        file_name = Path(file_path).name
                        message = f"Error: {file_name}: {data}"
                        print(f"✗ {message}")

                    if progress_callback:
                        progress_callback(indexed_count + error_count, total, message)

        except Exception as e:
            print(f"\n✗ Parallel processing error: {str(e)}")
            print("Falling back to sequential mode...")
            # Fall back to sequential if parallel fails
            remaining_files = files_to_index[indexed_count + error_count:]
            seq_indexed, seq_errors = self._scan_sequential(remaining_files, progress_callback)
            indexed_count += seq_indexed
            error_count += seq_errors

        return indexed_count, error_count

    def get_stats(self):
        """Get database statistics"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute('SELECT COUNT(*) FROM documents')
        total_docs = cursor.fetchone()[0]

        cursor.execute('SELECT file_type, COUNT(*) FROM documents GROUP BY file_type')
        by_type = cursor.fetchall()

        conn.close()

        print("\nDatabase Statistics:")
        print(f"Total documents: {total_docs}")
        print("\nBy file type:")
        for file_type, count in by_type:
            print(f"  {file_type}: {count}")


if __name__ == '__main__':
    # Required for multiprocessing on Windows
    multiprocessing.freeze_support()

    print("=" * 80)
    print("Document Indexer")
    print("=" * 80)

    indexer = DocumentIndexer()
    indexer.scan_and_index()  # Will use parallel processing by default
    indexer.get_stats()
